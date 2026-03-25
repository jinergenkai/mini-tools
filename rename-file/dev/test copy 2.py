import os
import re
from docx import Document
import pprint
import win32com.client  # Dùng để xử lý file .doc (Windows)

# Đọc danh sách từ khóa prefix từ file match.txt và ignore.txt
def load_keywords(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        keywords = [line.strip() for line in f.readlines()]
    return keywords

# Đọc file DOCX
def read_docx_text(filepath):
    doc = Document(filepath)
    return doc  # Trả về object Document để có thể lấy thêm thông tin về font size

# Đọc file DOC (cũ) - Windows
def read_doc_text(filepath):
    # Convert relative path to absolute path
    abs_path = os.path.abspath(filepath)
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Ẩn ứng dụng Word
    doc = word.Documents.Open(abs_path)
    full_text = doc.Content.Text  # Lấy toàn bộ nội dung văn bản
    doc.Close()
    word.Quit()
    return full_text

import os

def read_doc_paragraphs(filepath):
    try:
        # Cố gắng tạo đối tượng Word nếu có Office
        abs_path = os.path.abspath(filepath)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path)
        
        paragraphs = []
        
        # Đọc từng đoạn văn bản
        for p in doc.Paragraphs:
            text = p.Range.Text.strip()
            if text:  # Chỉ lấy các đoạn không rỗng
                # Kiểm tra căn chỉnh đoạn văn (Alignment)
                is_centered = (p.Alignment == 1)  # 1 tương ứng với căn giữa (wdAlignParagraphCenter)
                paragraphs.append({
                    'text': text,
                    'is_centered': is_centered
                })
        
        doc.Close()
        word.Quit()
        
        return paragraphs
    
    except ImportError:
        print("❌ Không thể import win32com.client. Đảm bảo rằng Microsoft Office đã được cài đặt.")
    
    except Exception as e:
        print(f"❌ Lỗi khi mở file .doc: {e}")

# Lấy font size lớn nhất từ dòng (paragraph)
def get_max_font_size(paragraph):
    max_font_size = 0
    for run in paragraph.runs:
        if run.font.size:  # Nếu có thông tin về kích thước font
            max_font_size = max(max_font_size, run.font.size.pt)  # Lấy giá trị lớn nhất (pt)
    return max_font_size

# Xử lý văn bản, lọc theo các quy tắc
def process_text(doc, match_keywords, ignore_keywords):
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title_line = None
    lines_data = []
    max_font_size_overall = 0  # Lưu lại font size lớn nhất toàn bộ văn bản

    # 1. Tính toán font size lớn nhất của toàn bộ văn bản
    for paragraph in doc.paragraphs:
        max_font_size_overall = max(max_font_size_overall, get_max_font_size(paragraph))

    # 2. Xử lý từng dòng để đánh điểm
    for line in lines[:20]:  # Chỉ xử lý 20 dòng đầu tiên
        line_data = {
            "text": line.strip(),
            "points": 0,
            "is_uppercase": line.isupper(),
            "is_centered": False,  # Cần kiểm tra xem dòng có căn giữa không
            "font_size": None,  # Giả sử không có thông tin font size trực tiếp trong .docx
        }

        # Tìm font size của dòng (paragraph) chứa dòng hiện tại
        max_font_size = 0
        for paragraph in doc.paragraphs:
            if line.strip() in paragraph.text:  # Kiểm tra nếu dòng có trong paragraph
                max_font_size = get_max_font_size(paragraph)
                line_data["font_size"] = max_font_size
                break

        # Đánh giá căn giữa
        if line_data["text"].strip() == line_data["text"].center(len(line_data["text"].strip())):
            line_data["is_centered"] = True
            line_data["points"] += 1

        # 3. Kiểm tra từ khóa match (trong match.txt) và ignore (trong ignore.txt)
        for keyword in match_keywords:
            if keyword in line_data["text"]:
                line_data["points"] += 1  # +1 điểm nếu có match từ khóa
        for keyword in ignore_keywords:
            if keyword in line_data["text"]:
                line_data["points"] -= 1  # -1 điểm nếu có match từ khóa ignore
        
        # 4. Kiểm tra dòng có phải là chữ hoa hết không
        if line_data["is_uppercase"]:
            line_data["points"] += 1  # +1 điểm nếu dòng viết hoa hết
        
        # 5. Kiểm tra dòng có phải là dòng có font size lớn nhất
        if max_font_size == max_font_size_overall:
            line_data["points"] += 1  # +1 điểm nếu dòng có font size lớn nhất toàn bộ văn bản

        # Lưu dữ liệu dòng để xét chọn tiêu đề
        lines_data.append(line_data)

    # 6. Chọn dòng có điểm số cao nhất
    title_line = max(lines_data, key=lambda x: x["points"])
    title = title_line["text"]

    pprint.pprint(lines_data)

    # 7. Thêm năm vào cuối nếu tìm thấy
    year_match = re.search(r"\b(20\d{2})\b", "\n".join([p.text for p in doc.paragraphs]))  # Cập nhật ở đây
    if year_match:
        title += f" {year_match.group(1)}"
    
    # 8. Xử lý chiều dài tên file, không quá 200 ký tự
    filename = title.strip()
    filename = filename[:200]  # Giới hạn 200 ký tự
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)  # Lọc các ký tự không hợp lệ trong tên file
    
    return filename

# def clean_text(text):
#     # Loại bỏ ký tự điều khiển và null byte (trừ newline \n)
#     return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)

def clean_text(text):
    # Loại bỏ các control character (trừ newline và tab)
    return re.sub(r'[^\x20-\x7E\n\t]', '', text)



# Hàm đổi tên file
def clean_text(text):
    # Remove control characters but keep newlines and tabs
    return ''.join(char for char in text if char == '\n' or char == '\t' or (ord(char) >= 32))

def rename_file_with_rules(filepath, match_keywords, ignore_keywords):
    file_extension = os.path.splitext(filepath)[1].lower()  # Lấy phần mở rộng của file (.doc hoặc .docx)
    
    if file_extension == ".docx":
        doc = read_docx_text(filepath)

    elif file_extension == ".doc":
        paragraphs = read_doc_paragraphs(filepath)  # ✅ chuẩn từng đoạn
        print(f"Đọc file .doc thành công: {filepath}")
        print(f"Đọc {len(paragraphs)} đoạn văn bản.")
        for i, para in enumerate(paragraphs):
            print(f"Đoạn {i+1}: {para}")

        # return 
        doc = Document()
        for para in paragraphs:
            safe_para = clean_text(para)
            doc.add_paragraph(safe_para)

    
    else:
        raise ValueError("Unsupported file format. Only .doc and .docx are supported.")
    
    
    new_filename = process_text(doc, match_keywords, ignore_keywords)
    new_path = os.path.join(os.path.dirname(filepath), new_filename + file_extension)
    
    # os.rename(filepath, new_path)
    print(f"✅ Đổi tên file thành: {new_filename}{file_extension}")

# 🧪 Thử demo
match_keywords = load_keywords("match.txt")  # Load từ khóa từ file match.txt
ignore_keywords = load_keywords("ignore.txt")  # Load từ khóa cần loại bỏ từ ignore.txt
rename_file_with_rules("123.docx", match_keywords, ignore_keywords)  # Đổi tên file .docx
rename_file_with_rules("111.doc", match_keywords, ignore_keywords)   # Đổi tên file .doc
