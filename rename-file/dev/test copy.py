import os
import re
from docx import Document

# Đọc danh sách từ khóa prefix từ file match.txt
def load_keywords(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        keywords = [line.strip() for line in f.readlines()]
    return keywords

# Đọc danh sách từ khóa cần bỏ qua từ file ignore.txt
def load_ignore_keywords(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        ignore_keywords = [line.strip() for line in f.readlines()]
    return ignore_keywords

# Đánh điểm các dòng theo các tiêu chí
def score_line(line, match_keywords, ignore_keywords, font_sizes):
    score = 0
    # 1. Nếu dòng có từ khóa trùng với match.txt, cộng 1 điểm
    if any(keyword.lower() in line.lower() for keyword in match_keywords):
        score += 1
    
    # 2. Nếu dòng là căn giữa, cộng 1 điểm
    if line == line.strip() and line.strip() == line.strip().center(len(line)):
        score += 1
    
    # 3. Nếu dòng viết hoa hết, cộng 1 điểm
    if line.isupper():
        score += 1
    
    # 4. Nếu dòng có font chữ lớn nhất trong 20 dòng đầu tiên, cộng 1 điểm
    if font_sizes and font_sizes.get(line, 0) == max(font_sizes.values()):
        score += 1
    
    # 5. Nếu dòng chứa từ khóa trong ignore.txt, trừ 1 điểm
    if any(ignore.lower() in line.lower() for ignore in ignore_keywords):
        score -= 1
    
    return score

# Xử lý văn bản và tìm tiêu đề
def process_text(doc, text, match_keywords, ignore_keywords):
    lines = text.split("\n")
    
    # Lưu trữ các thuộc tính của các dòng
    font_sizes = {}
    for para in doc.paragraphs[:20]:
        for run in para.runs:
            line = para.text.strip()
            if line:  # Nếu đoạn không rỗng
                font_sizes[line] = max(font_sizes.get(line, 0), run.font.size.pt if run.font.size else 0)
    
    # Đánh điểm cho từng dòng
    scored_lines = []
    for line in lines[:20]:
        score = score_line(line, match_keywords, ignore_keywords, font_sizes)
        scored_lines.append((line, score))
    
    # Tìm dòng có điểm cao nhất làm tiêu đề
    best_title = max(scored_lines, key=lambda x: x[1])[0]
    
    return best_title.strip()

# Đọc file DOCX
def read_docx_text(filepath):
    doc = Document(filepath)
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    return doc, full_text

# Hàm đổi tên file
def rename_file_with_rules(filepath, match_keywords, ignore_keywords):
    doc, original_text = read_docx_text(filepath)
    new_title = process_text(doc, original_text, match_keywords, ignore_keywords)
    
    # Tạo tên file từ tiêu đề đã tìm được
    new_filename = new_title[:200]  # Giới hạn 200 ký tự
    new_filename = re.sub(r'[<>:"/\\|?*]', '', new_filename)  # Lọc các ký tự không hợp lệ trong tên file
    # new_path = os.path.join(os.path.dirname(filepath), new_filename + ".docx")
    
    # os.rename(filepath, new_path)
    print(f"✅ Đổi tên file thành: {new_filename}.docx")

# 🧪 Thử demo
match_keywords = load_keywords("match.txt")  # Load từ khóa từ file match.txt
ignore_keywords = load_ignore_keywords("ignore.txt")  # Load từ khóa cần bỏ qua từ file ignore.txt
rename_file_with_rules("123.docx", match_keywords, ignore_keywords)
